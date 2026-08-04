"""Convert the pipeline's illustrated Markdown into a polished DOCX.

Design system: narrative_proposal preset + editorial_cover first page.
The implementation intentionally uses native Word headings, numbering, tables,
images, links, headers, and page-number fields instead of flattening Markdown.
"""

from __future__ import annotations

import re
import zipfile
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from PIL import Image


PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "203748"
MUTED = "626B75"
LIGHT_FILL = "F4F6F9"
TABLE_HEADER_FILL = "F4F6F9"
BORDER = "CAD2DC"
WHITE = "FFFFFF"
LINK_BLUE = "0563C1"


def _hex(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 ascii_font: str = "Calibri", east_asia_font: str = "Microsoft YaHei") -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _hex(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_margins(cell, top: int = CELL_TOP_BOTTOM_DXA, start: int = CELL_SIDE_DXA,
                      bottom: int = CELL_TOP_BOTTOM_DXA, end: int = CELL_SIDE_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{side}"))
        if tag is None:
            tag = OxmlElement(f"w:{side}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def _shade(element, fill: str) -> None:
    props = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def _paragraph_border(paragraph, *, side: str = "left", color: str = BLUE,
                      size: int = 18, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def _paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _keep(paragraph, *, with_next: bool = False, together: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = with_next
    paragraph.paragraph_format.keep_together = together


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])
    set_run_font(run, size=9, color=MUTED)


def _add_hyperlink(paragraph, text: str, url: str):
    relation = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(`[^`]+`|\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|\[[^\]]+\]\([^)]+\))"
)


MACRO_TEXT = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ", "epsilon": "ε",
    "varepsilon": "ε", "zeta": "ζ", "eta": "η", "theta": "θ", "iota": "ι",
    "kappa": "κ", "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ", "pi": "π",
    "rho": "ρ", "sigma": "σ", "tau": "τ", "upsilon": "υ", "phi": "φ",
    "varphi": "φ", "chi": "χ", "psi": "ψ", "omega": "ω",
    "Gamma": "Γ", "Delta": "Δ", "Theta": "Θ", "Lambda": "Λ", "Xi": "Ξ",
    "Pi": "Π", "Sigma": "Σ", "Phi": "Φ", "Psi": "Ψ", "Omega": "Ω",
    "mid": "|", "in": "∈", "notin": "∉", "infty": "∞", "partial": "∂",
    "nabla": "∇", "pm": "±", "mp": "∓", "neq": "≠", "equiv": "≡",
    "propto": "∝", "sim": "~", "to": "→", "rightarrow": "→",
    "leftarrow": "←", "Rightarrow": "⇒", "langle": "⟨", "rangle": "⟩",
    "ldots": "…", "dots": "…", "cdots": "…", "sum": "Σ", "prod": "Π",
    "int": "∫", "sqrt": "√", "forall": "∀", "exists": "∃",
    "bigl": "", "bigr": "", "Bigl": "", "Bigr": "", "biggl": "", "biggr": "",
    "quad": " ", "qquad": "  ", "displaystyle": "",
    # Standard operator names render as themselves, minus the backslash.
    **{name: name for name in (
        "log", "ln", "exp", "min", "max", "arg", "det", "dim", "gcd",
        "sin", "cos", "tan", "lim", "sup", "inf", "deg", "ker",
    )},
}
_MACRO_PATTERN = re.compile(r"\\([A-Za-z]+)")
_OPEN_BRACE, _CLOSE_BRACE = "\x01", "\x02"


def latex_to_plain(value: str) -> str:
    """Convert the small LaTeX subset emitted by the writing pipeline to readable text."""
    value = value.strip()
    for opening, closing in ((r"\[", r"\]"), ("$$", "$$"), (r"\(", r"\)")):
        if value.startswith(opening) and value.endswith(closing):
            value = value[len(opening):-len(closing)].strip()
    # Escaped braces are literal characters, not grouping. Park them so the
    # grouping-aware passes below see only real LaTeX groups.
    value = value.replace(r"\{", _OPEN_BRACE).replace(r"\}", _CLOSE_BRACE)
    while True:
        unwrapped = re.sub(r"\\(?:text|operatorname|mathrm|mathbf)\{([^{}]*)\}", r"\1", value)
        if unwrapped == value:
            break
        value = unwrapped
    value = value.replace(r"\_", "_")
    # Collapse sub/superscripts repeatedly: one pass leaves the outer group of a
    # nested script (p_{\theta_{old}}) intact, which then defeats \frac matching.
    while True:
        scripted = re.sub(r"_\{([^{}]*)\}", r"_\1", value)
        scripted = re.sub(r"\^\{([^{}]*)\}", r"^\1", scripted)
        if scripted == value:
            break
        value = scripted
    fraction = re.compile(r"\\frac\{([^{}]*)\}\{([^{}]*)\}")
    while fraction.search(value):
        value = fraction.sub(lambda match: f"({match.group(1)}) / ({match.group(2)})", value)
    value = value.replace("{,}", ",")
    value = value.replace(r"\times", " × ").replace(r"\cdot", " · ")
    value = value.replace(r"\approx", "≈").replace(r"\le", "≤").replace(r"\ge", "≥")
    value = value.replace(r"\;", " ").replace(r"\,", " ").replace(r"\ ", " ")
    value = value.replace(r"\!", "").replace(r"\left", "").replace(r"\right", "")
    value = _MACRO_PATTERN.sub(
        lambda match: MACRO_TEXT.get(match.group(1), match.group(0)), value)
    value = value.replace("{", "").replace("}", "")
    value = value.replace(_OPEN_BRACE, "{").replace(_CLOSE_BRACE, "}")
    return re.sub(r"\s+", " ", value).strip()


def inline_math_to_plain(value: str) -> str:
    value = re.sub(r"\\\((.+?)\\\)", lambda match: latex_to_plain(match.group(1)), value)
    return re.sub(r"(?<!\\)\$([^$\n]+)\$", lambda match: latex_to_plain(match.group(1)), value)


def add_inline(paragraph, text: str, *, default_size: float | None = None,
               default_color: str | None = None) -> None:
    text = inline_math_to_plain(text)

    def apply_default(run, *, bold: bool | None = None, italic: bool | None = None) -> None:
        if default_size is not None or default_color is not None:
            set_run_font(run, size=default_size, color=default_color, bold=bold, italic=italic)
        else:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic

    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            apply_default(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            apply_default(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=9.5, color=DARK_BLUE, ascii_font="Consolas",
                         east_asia_font="Microsoft YaHei")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "EEF2F6")
            run._element.get_or_add_rPr().append(shd)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            apply_default(run, italic=True)
        else:
            link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                _add_hyperlink(paragraph, link.group(1), link.group(2))
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        apply_default(run)


def _make_numbering(document: Document, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "279")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.extend([start, fmt, text, jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.append(num_pr)


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _table_widths(rows: list[list[str]]) -> list[int]:
    columns = max(len(row) for row in rows)
    scores = []
    for index in range(columns):
        values = [re.sub(r"[*`]", "", row[index]) if index < len(row) else "" for row in rows]
        scores.append(max(4, min(42, max((len(value) for value in values), default=4))))
    minimum = 1100 if columns <= 4 else 720
    available = CONTENT_WIDTH_DXA - minimum * columns
    score_total = sum(scores)
    widths = [minimum + round(available * score / score_total) for score in scores]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    widths = _table_widths(rows)
    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            value = values[column_index] if column_index < len(values) else ""
            add_inline(paragraph, value, default_size=9.4)
            if row_index == 0:
                _shade(cell._tc, TABLE_HEADER_FILL)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = _hex(INK_BLUE)
            if column_index > 0 and len(value) < 18:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_props = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_props.append(repeat)
    _set_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _configure_styles(document: Document) -> tuple[int, int]:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = _hex("202124")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _hex(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = _hex(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)

    if "Lead" not in styles:
        lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Lead"]
    lead.font.name = "Calibri"
    lead.font.size = Pt(12)
    lead.font.color.rgb = _hex(INK_BLUE)
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    lead.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lead.paragraph_format.space_before = Pt(2)
    lead.paragraph_format.space_after = Pt(12)
    lead.paragraph_format.line_spacing = 1.3

    if "Formula" not in styles:
        formula = styles.add_style("Formula", WD_STYLE_TYPE.PARAGRAPH)
    else:
        formula = styles["Formula"]
    formula.font.name = "Calibri"
    formula.font.size = Pt(10.5)
    formula.font.bold = True
    formula.font.color.rgb = _hex(INK_BLUE)
    formula._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(5)
    formula.paragraph_format.space_after = Pt(10)
    formula.paragraph_format.keep_together = True

    return _make_numbering(document, "bullet"), _make_numbering(document, "decimal")


def _configure_section(document: Document, title: str, language: str) -> None:
    section = document.sections[0]
    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(PAGE_HEIGHT_DXA)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = paragraph.add_run(title[:52])
    set_run_font(left, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_label = "Technical Deep Dive" if language == "en" else "技术深度解析"
    label = footer_p.add_run(f"{footer_label}  ·  ")
    set_run_font(label, size=9, color=MUTED)
    _add_page_field(footer_p)


def _extract_title(lines: list[str]) -> tuple[str, str, list[str]]:
    title = "技术博客"
    subtitle = ""
    remaining = lines[:]
    while remaining and not remaining[0].strip():
        remaining.pop(0)
    if remaining and re.match(r"^#\s+", remaining[0]):
        title = re.sub(r"^#\s+", "", remaining.pop(0)).strip()
    while remaining and not remaining[0].strip():
        remaining.pop(0)
    if remaining and remaining[0].startswith(">"):
        subtitle = remaining.pop(0).lstrip("> ").strip()
    return title, subtitle, remaining


def _add_cover(document: Document, title: str, subtitle: str, source_label: str | None,
               language: str) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(118)
    spacer.paragraph_format.space_after = Pt(0)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("TECHNICAL DEEP DIVE")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(10)
    title_p.paragraph_format.keep_together = True
    run = title_p.add_run(title)
    set_run_font(run, size=28, color=INK_BLUE, bold=True)

    if subtitle:
        subtitle_p = document.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_p.paragraph_format.space_after = Pt(28)
        run = subtitle_p.add_run(subtitle)
        set_run_font(run, size=14, color=DARK_BLUE)

    descriptor = document.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_before = Pt(42)
    descriptor.paragraph_format.space_after = Pt(8)
    descriptor_text = (
        "Synthesized from a technical talk and its accompanying slides"
        if language == "en" else "基于技术演讲与配套材料整理"
    )
    run = descriptor.add_run(descriptor_text)
    set_run_font(run, size=10.5, color=MUTED, italic=True)
    if source_label:
        source_p = document.add_paragraph()
        source_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_p.paragraph_format.space_after = Pt(4)
        run = source_p.add_run(source_label)
        set_run_font(run, size=9.5, color=MUTED)

    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(34)
    run = date_p.add_run(date.today().isoformat())
    set_run_font(run, size=10, color=MUTED)
    date_p.add_run().add_break(WD_BREAK.PAGE)


def _image_size(path: Path) -> tuple[float, float]:
    with Image.open(path) as image:
        width, height = image.size
    max_width = 6.5
    max_height = 5.7
    ratio = min(max_width / width, max_height / height)
    return width * ratio, height * ratio


def _add_image(document: Document, path: Path, alt: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    _keep(paragraph, with_next=True, together=True)
    width, height = _image_size(path)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width), height=Inches(height))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    doc_pr.set("title", alt[:120])


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        # Be lenient with model-generated Markdown. Some translations emit a
        # single dash in a narrow alignment cell (for example ``|:-:|``),
        # which is semantically still a table separator even though CommonMark
        # normally expects at least three dashes.
        if not all(re.fullmatch(r":?-+:?", value.replace(" ", "")) for value in values):
            rows.append(values)
        index += 1
    return rows, index


def markdown_to_docx(markdown_path: Path, output_path: Path, *, source_label: str | None = None,
                     language: str = "zh") -> Path:
    if language not in {"zh", "en"}:
        raise ValueError("language must be 'zh' or 'en'")
    markdown_path = markdown_path.resolve()
    output_path = output_path.resolve()
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    title, subtitle, lines = _extract_title(lines)

    document = Document()
    bullet_num, decimal_num = _configure_styles(document)
    _configure_section(document, title, language)
    _add_cover(document, title, subtitle, source_label, language)

    index = 0
    body_paragraph_count = 0
    previous_list_kind: str | None = None
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped in {r"\[", "$$"}:
            closing = r"\]" if stripped == r"\[" else "$$"
            formula_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != closing:
                formula_lines.append(lines[index].strip())
                index += 1
            if index < len(lines):
                index += 1
            paragraph = document.add_paragraph(style="Formula")
            _paragraph_shading(paragraph, LIGHT_FILL)
            _paragraph_border(paragraph, side="bottom", color=BORDER, size=6, space=4)
            add_inline(paragraph, latex_to_plain(" ".join(formula_lines)))
            previous_list_kind = None
            continue
        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            image_path = (markdown_path.parent / image_match.group(2)).resolve()
            if image_path.is_file():
                _add_image(document, image_path, image_match.group(1))
            else:
                warning = document.add_paragraph()
                warning.style = "Caption"
                add_inline(warning, f"[图片缺失：{image_match.group(2)}]", default_size=9,
                           default_color="9B1C1C")
            index += 1
            previous_list_kind = None
            continue
        if stripped.startswith("|"):
            rows, index = _parse_table(lines, index)
            _add_table(document, rows)
            previous_list_kind = None
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = min(3, len(heading.group(1)) - 1)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, heading.group(2))
            index += 1
            previous_list_kind = None
            continue
        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.line_spacing = 1.25
            _paragraph_border(paragraph)
            _paragraph_shading(paragraph, LIGHT_FILL)
            add_inline(paragraph, stripped.lstrip("> "), default_size=11.5, default_color=INK_BLUE)
            index += 1
            previous_list_kind = None
            continue
        bullet = re.match(r"^[-+*]\s+(.+)$", stripped)
        number = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if bullet or number:
            list_kind = "bullet" if bullet else "number"
            if number and previous_list_kind != "number":
                decimal_num = _make_numbering(document, "decimal")
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.208
            _apply_numbering(paragraph, bullet_num if bullet else decimal_num)
            add_inline(paragraph, (bullet or number).group(1))
            index += 1
            previous_list_kind = list_kind
            continue
        if re.match(r"^\*图(?:\s*\d+)?[：:].+\*$", stripped):
            paragraph = document.add_paragraph(style="Caption")
            add_inline(paragraph, stripped.strip("*"))
            index += 1
            previous_list_kind = None
            continue

        paragraph = document.add_paragraph(style="Lead" if body_paragraph_count == 0 else "Normal")
        add_inline(paragraph, stripped)
        body_paragraph_count += 1
        index += 1
        previous_list_kind = None

    properties = document.core_properties
    properties.title = title
    properties.subject = subtitle or (
        "Technical blog synthesized from a video and slide deck"
        if language == "en" else "技术视频与 PPT 整理的技术博客"
    )
    properties.author = "Technical Blog Pipeline"
    properties.keywords = "technical blog, video, PPT, AI systems"
    properties.comments = "Generated from user-provided media and slides."
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def audit_docx(path: Path) -> dict[str, object]:
    """Run deterministic integrity, resource, numbering, and table-geometry checks."""
    path = path.resolve()
    document = Document(path)
    with zipfile.ZipFile(path) as archive:
        zip_error = archive.testzip()
        media_parts = [name for name in archive.namelist() if name.startswith("word/media/")]

    geometry_errors: list[str] = []
    for table_index, table in enumerate(document.tables, 1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        declared = int(tbl_w.get(qn("w:w"), "0")) if tbl_w is not None else 0
        grid = [int(node.get(qn("w:w"), "0")) for node in table._tbl.tblGrid]
        if declared != CONTENT_WIDTH_DXA or sum(grid) != declared:
            geometry_errors.append(f"table {table_index}: tblW/grid mismatch")
        for row_index, row in enumerate(table.rows, 1):
            widths = []
            for cell in row.cells:
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                widths.append(int(tc_w.get(qn("w:w"), "0")) if tc_w is not None else 0)
            if widths != grid:
                geometry_errors.append(f"table {table_index} row {row_index}: tcW mismatch")

    missing_alt = []
    for index, shape in enumerate(document.inline_shapes, 1):
        if not (shape._inline.docPr.get("descr") or "").strip():
            missing_alt.append(index)
    numbered = sum(1 for paragraph in document.paragraphs
                   if paragraph._p.xpath("./w:pPr/w:numPr"))
    headings = sum(1 for paragraph in document.paragraphs
                   if paragraph.style.name.startswith("Heading"))
    passed = zip_error is None and not geometry_errors and not missing_alt and len(media_parts) == len(document.inline_shapes)
    return {
        "path": str(path), "size": path.stat().st_size, "pass": passed,
        "paragraphs": len(document.paragraphs), "headings": headings,
        "numbered_paragraphs": numbered, "tables": len(document.tables),
        "embedded_images": len(document.inline_shapes), "media_parts": len(media_parts),
        "zip_integrity": "pass" if zip_error is None else f"fail:{zip_error}",
        "table_geometry": "pass" if not geometry_errors else geometry_errors,
        "images_missing_alt": missing_alt,
        "visual_render_qa": "not_run_by_pipeline",
    }


if __name__ == "__main__":
    import argparse

    cli = argparse.ArgumentParser(description="把图文 Markdown 转为排版好的 DOCX")
    cli.add_argument("markdown", type=Path)
    cli.add_argument("output", type=Path)
    cli.add_argument("--source-label")
    cli.add_argument("--language", choices=["zh", "en"], default="zh")
    values = cli.parse_args()
    markdown_to_docx(values.markdown, values.output, source_label=values.source_label,
                     language=values.language)
